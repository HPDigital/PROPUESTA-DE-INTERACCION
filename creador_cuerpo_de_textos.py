
#!/usr/bin/env python
# coding: utf-8

# In[2]:


import os
from typing import Any, Dict, Optional
from docx import Document
from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
from langchain_openai.chat_models import ChatOpenAI
from PyPDF2 import PdfReader

# ------------------------------------------------------------------------------
# CONFIGURACIÓN E INICIALIZACIÓN
# ------------------------------------------------------------------------------

# Cargar variables de entorno desde el archivo .env
load_dotenv()

# Obtener la API key de las variables de entorno y validarla
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("La clave de la API de OpenAI no se encontró en las variables de entorno.")

# Crear el cliente LLM
llm = ChatOpenAI(api_key=OPENAI_API_KEY, model="o3-mini-2025-01-31", temperature=1)

# ------------------------------------------------------------------------------
# DEFINICIÓN DE PROMPTS
# ------------------------------------------------------------------------------

# Prompt para generar un resumen académico a partir de una sección
PROMPT_RESUMEN = PromptTemplate(
    input_variables=["titulo", "subtitulo", "contenido"],
    template="""
A continuación se muestra un ejemplo de prompt completo que puedes utilizar para que el modelo busque, recopile y desarrolle información pertinente y detallada a partir del título y subtítulo proporcionados:

---

**Prompt:**

Eres un experto académico en el tema relacionado con el título, subtítulo y contenido que se te proporcionan. 
Tu tarea es buscar y recopilar información precisa, detallada y relevante, integrando definiciones, antecedentes, análisis crítico y ejemplos ilustrativos que enriquezcan la comprensión del tema. Para ello, sigue estas instrucciones:

1. Contextualización y Definiciones:
   - Explica de manera clara y precisa los conceptos clave relacionados con el *TÍTULO* y el *SUBTÍTULO*.  
   - Define términos técnicos y contextualiza el tema dentro de su ámbito académico y/o profesional.

2. Respuestas directas y sin introduciones:
   - Brinda respuestas directas y clras sin rpetir el texto del titulo y subtitulo ni el titulo.  
   - Proporciona la respuestas sin introducciones ni bla bla

3. Análisis y Desarrollo del Tema: 
   - Desglosa los aspectos más importantes y pertinentes que abarca el *SUBTÍTULO*, resaltando los puntos críticos y sus implicaciones.  
   - Incluye ejemplos, casos de estudio o aplicaciones prácticas que permitan comprender mejor el impacto y la relevancia del tema.  
   - Analiza las ventajas, desafíos y perspectivas futuras, proporcionando un enfoque crítico y reflexivo.

4. Estructura y Redacción: 
   - Organiza tu respuesta en secciones claramente diferenciadas: *Introducción*, *Desarrollo* y *Conclusión*.  
   - Emplea un lenguaje formal, académico y coherente, cuidando la precisión en el uso de términos y la claridad en la exposición de ideas.  
   - Asegúrate de que la respuesta sea extensa y detallada, evitando generalizaciones y proporcionando información específica y contrastada.

5. Formato de Respuesta: 
   - La respuesta final debe ser un único texto fluido, dividido en párrafos bien estructurados, que integre de forma natural todos los puntos solicitados.  
   - Si es pertinente, menciona las fuentes o referencias clave (sin necesidad de incluir citas literales) que respaldan la información ofrecida.


El texto debe contener explicaiones acdemicas con los siguientes detalles:

Título: {titulo}
SubTítulo: {subtitulo}
Contenido: {contenido}

Proporciona una respuesta amplia y detallado.
Deber ir que vaya acorde con una PROPUESTA DE INTERACCIÓN ACADÉMICA ENTRE LA UNIVERSIDAD CATÓLICA BOLIVIANA Y EL SECTOR EMPRESARIAL DE COCHABAMBA."""
)

# Prompt para adaptar el contenido a un estilo basado en un ejemplo extraído de un PDF
PROMPT_ESTILO = PromptTemplate(
    input_variables=["texto_estilo", "contenido"],
    template="""Utiliza el siguiente ejemplo de redacción 100% humana obtenido de un documento PDF para aprender a redactar respuestas largas, fluidas y con un uso adecuado de los conectores en español.  
No se trata de resumir el contenido del ejemplo, sino de imitar su estilo gramatical de redaccion.

Siguiendo las siguientes reglas:
1. El largo de cada oracion debe ser mayor que lo habitual.
2. Toma el texto y divídelo en dos o mas frases completas, asegurándote de mantener la esencia y el significado original sin cambiar demasiado las palabras. 
3. Cada frase debe ser clara y estructurada, respetando la coherencia y fluidez del contenido.
4. Mantén el tono formal y académico. 
5. No elimines información clave, pero reorganiza el contenido si es necesario para que cada frase tenga sentido por sí misma
6. Utiliza menos puntos separadores y conecta las oraciones usando conectores en español.
7. Evita redundancias y divagaciones; expón las ideas de forma directa.
8. Usa un lenguaje accesible, priorizando palabras comunes sobre términos complejos.
9. Adapta el nivel a una audiencia universitaria, sin tecnicismos excesivos.
10. Las frases deben ser largas y contener, al menos, dos oraciones conectadas.
11. Conecta las ideas de forma natural, sin transiciones forzadas.
12. Respeta las normas convencionales de puntuación.
13. Varía la estructura sintáctica para evitar monotonía.
14. La respuesta debe ser una sola frase compuesta de varias oraciones.
15. Evita conectores innecesarios (por ejemplo: "de hecho", "además", "por lo tanto", etc.).
16. Asegúrate de que las oraciones sean coherentes y fáciles de entender.
17. Considera el tono y el contexto original.
18. Tu respuesta debe estar estrictamente relacionada con el subtítulo proporcionado, explicando de forma directa y precisa los aspectos que se mencionan en él.
19. No te desvíes del tema central del subtítulo y relaciona el contenido exclusivamente con esa temática.
20.  La respuesta final debe tener oraciones largas que conformen desde un minimo dos frases hasta un maximo de cinco frases. 
21. Debes esoger la cantidad de frases que contendra la respuesta de forma aleatoria, entre minimo dos frases hasta maximo cinco frases.

Basándote en las instrucciones, y tomando en cuenta que estamos escribiendo un libro academico sobre la interaccion acdemica-empresarial.
Este texto deb detllar la inetraccion eentre la carrrade de adminitracion de empresas de la universidad catolica boliviana y el conh¿junto de empresas de cochabamba.

Output Format:
El texto debe ser FORMAL ACADEMICO, imitando el estilo gramatical de redacción y el uso de palabras conectaoras del idioma español. 
El texto debe ser formal y academico.


Ejemplo de estilo:
{texto_estilo}

Ahora, redacta el siguiente contenido usando un estilo similar:
{contenido}

Respuesta:"""
)

# ------------------------------------------------------------------------------
# FUNCIONES AUXILIARES
# ------------------------------------------------------------------------------

def invoke_prompt(prompt: PromptTemplate, inputs: Dict[str, str]) -> str:
    """
    Invoca el prompt dado junto al LLM y retorna la respuesta procesada.
    """
    chain = prompt | llm
    response = chain.invoke(inputs)
    return response.content.strip()


def leer_pdf(ruta_pdf: str, max_tokens: int = 80000) -> str:
    """
    Lee y extrae el texto de un archivo PDF, retornando los primeros max_tokens tokens.

    Args:
        ruta_pdf (str): Ruta del archivo PDF.
        max_tokens (int, opcional): Número máximo de tokens a extraer. Por defecto 50000.

    Returns:
        str: Texto extraído del PDF.
    """
    try:
        reader = PdfReader(ruta_pdf)
        tokens_totales = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                tokens_page = page_text.split()  # Tokenización simple por espacios
                tokens_totales.extend(tokens_page)
                if len(tokens_totales) >= max_tokens:
                    tokens_totales = tokens_totales[:max_tokens]
                    break
        print(f"Total tokens leídos del PDF: {len(tokens_totales)}")
        return " ".join(tokens_totales)
    except Exception as e:
        print(f"Error al leer el PDF: {e}")
        return ""

# ------------------------------------------------------------------------------
# FUNCIONES DE INTERACCIÓN CON EL LLM
# ------------------------------------------------------------------------------

def interactuar_proposito(titulo: str, subtitulo: str, contenido: str) -> str:
    """
    Genera un resumen académico para una sección específica.

    Args:
        titulo (str): Título global del documento.
        subtitulo (str): Título o subtítulo de la sección.
        contenido (str): Contenido de la sección.

    Returns:
        str: Resumen generado.
    """
    inputs = {"titulo": titulo, "subtitulo": subtitulo, "contenido": contenido}
    return invoke_prompt(PROMPT_RESUMEN, inputs)


def interactuar_estilo(contenido: str, ruta_pdf_estilo: str) -> str:
    """
    Adapta el contenido proporcionado a un estilo de redacción basado en un ejemplo extraído de un PDF.

    Args:
        contenido (str): Contenido al que se aplicará el estilo.
        ruta_pdf_estilo (str): Ruta del archivo PDF que contiene el ejemplo de estilo.

    Returns:
        str: Contenido adaptado al estilo indicado. Si falla la lectura del PDF, retorna el contenido original.
    """
    estilo_texto = leer_pdf(ruta_pdf_estilo)
    if not estilo_texto:
        return contenido
    inputs = {"texto_estilo": estilo_texto, "contenido": contenido}
    return invoke_prompt(PROMPT_ESTILO, inputs)


def interactuar_gpt(titulo: str, subtitulo: str, contenido: str, ruta_pdf_estilo: Optional[str] = None) -> str:
    """
    Genera el resumen de la sección y, opcionalmente, aplica un estilo de redacción.

    Args:
        titulo (str): Título global del documento.
        subtitulo (str): Título o subtítulo de la sección.
        contenido (str): Contenido de la sección.
        ruta_pdf_estilo (str, opcional): Ruta del PDF para aplicar estilo. Por defecto None.

    Returns:
        str: Texto final procesado.
    """
    resultado = interactuar_proposito(titulo, subtitulo, contenido)
    if ruta_pdf_estilo:
        resultado = interactuar_estilo(resultado, ruta_pdf_estilo)
    return resultado

# ------------------------------------------------------------------------------
# FUNCIONES PARA GENERAR EL DOCUMENTO
# ------------------------------------------------------------------------------

def procesar_indice(
    indice: Dict[str, Any],
    titulo_global: str,
    ruta_pdf_estilo: Optional[str],
    doc: Document,
    nivel: int
) -> None:
    """
    Recorre recursivamente la estructura del índice para agregar encabezados y párrafos al documento Word.

    Args:
        indice (dict): Estructura del índice.
        titulo_global (str): Título global del documento.
        ruta_pdf_estilo (str, opcional): Ruta al PDF para aplicar estilo, si se desea.
        doc (Document): Objeto Document de python-docx.
        nivel (int): Nivel de encabezado (heading).
    """
    for key, value in indice.items():
        doc.add_heading(key, level=nivel)
        if isinstance(value, dict):
            procesar_indice(value, titulo_global, ruta_pdf_estilo, doc, nivel + 1)
        else:
            contenido = str(value)
            resumen = interactuar_gpt(titulo_global, key, contenido, ruta_pdf_estilo)
            doc.add_paragraph(resumen)


def generar_resumenes_desde_json(
    json_data: Dict[str, Any],
    path_out: str,
    ruta_pdf_estilo: Optional[str] = None
) -> None:
    """
    Procesa el JSON con la estructura del documento y genera un archivo Word respetando la jerarquía.

    La salida contiene:
      - Un encabezado global (nivel 1).
      - Un encabezado "ÍNDICE DE CONTENIDO" (nivel 2).
      - Cada sección del índice en el orden y nivel jerárquico indicado.

    Args:
        json_data (dict): Estructura con "titulo" e "indice".
        path_out (str): Ruta de salida para el documento Word.
        ruta_pdf_estilo (str, opcional): Ruta del PDF para aplicar estilo (opcional).
    """
    doc = Document()
    titulo_global = json_data.get("titulo", "Documento")
    doc.add_heading(titulo_global, level=1)
    doc.add_heading("ÍNDICE DE CONTENIDO", level=2)

    indice = json_data.get("indice", {})
    procesar_indice(indice, titulo_global, ruta_pdf_estilo, doc, nivel=2)

    doc.save(path_out)
    print(f"Documento completado y guardado en: {path_out}")

# ------------------------------------------------------------------------------
# BLOQUE PRINCIPAL DE EJECUCIÓN
# ------------------------------------------------------------------------------

if __name__ == "__main__":
    # Ejemplo de JSON que respeta la estructura del índice
    json_data = {
        "titulo": "PROPUESTA DE INTERACCIÓN ACADÉMICA ENTRE LA UNIVERSIDAD Y EL SECTOR EMPRESARIAL",
        "indice": {
            "1. Introducción": {
                "1.1 Presentación de la Carrera y su Alineación con el Desarrollo Empresarial": "Contenido de la sección 1.1",
                "1.2 Importancia de la Relación entre la Academia y la Empresa": "Contenido de la sección 1.2",
                "1.3 Metodología Utilizada para el Diagnóstico del Sector Empresarial": "Contenido de la sección 1.3"
            },
            "2. Justificación de la Propuesta": {
                "2.1 Importancia de la Vinculación Universidad-Empresa en el Contexto Actual": "Contenido de la sección 2.1",
                "2.2 Alineación con el Modelo Institucional de la UCB y la Misión Educativa": "Contenido de la sección 2.2",
                "2.3 Impacto en la Mejora de la Calidad Educativa y el Desarrollo Empresarial": "Contenido de la sección 2.3",
                "2.4 Beneficios en Términos de Innovación, Empleabilidad y Responsabilidad Social": "Contenido de la sección 2.4"
            },
            "3. Objetivos de la Propuesta": {
                "3.1 Objetivo General": "Contenido de la sección 3.1",
                "3.2 Objetivos Específicos": {
                    "3.2.1 Fortalecer los Programas de Pasantías y Prácticas Empresariales": "Contenido de la sección 3.2.1",
                    "3.2.2 Fomentar la Investigación Aplicada y el Desarrollo Empresarial": "Contenido de la sección 3.2.2",
                    "3.2.3 Desarrollar Programas de Formación Continua y Capacitación Empresarial": "Contenido de la sección 3.2.3",
                    "3.2.4 Crear Espacios de Networking y Colaboración Academia-Empresa": "Contenido de la sección 3.2.4",
                    "3.2.5 Impulsar el Emprendimiento e Innovación Social": "Contenido de la sección 3.2.5",
                    "3.2.6 Facilitar la Transferencia de Tecnología y Conocimiento": "Contenido de la sección 3.2.6",
                    "3.2.7 Asegurar la Sostenibilidad del Modelo de Interacción": "Contenido de la sección 3.2.7"
                }
            },
            "4. Diagnóstico del Sector Empresarial en Cochabamba": {
                "4.1 Análisis PESTEL del Entorno Empresarial": {
                    "4.1.1 Factores Políticos": "Contenido de la sección 4.1.1",
                    "4.1.2 Factores Económicos": "Contenido de la sección 4.1.2",
                    "4.1.3 Factores Socioculturales": "Contenido de la sección 4.1.3",
                    "4.1.4 Factores Tecnológicos": "Contenido de la sección 4.1.4",
                    "4.1.5 Factores Ecológicos": "Contenido de la sección 4.1.5",
                    "4.1.6 Factores Legales": "Contenido de la sección 4.1.6"
                },
                "4.2 Impacto de la Transformación Digital en el Sector Empresarial": {
                    "4.2.1 Adopción de Tecnologías en las Empresas Cochabambinas": "Contenido de la sección 4.2.1",
                    "4.2.2 Digitalización de Procesos Empresariales y Tendencias en Automatización": "Contenido de la sección 4.2.2",
                    "4.2.3 Uso de Big Data, Inteligencia Artificial y Analítica Avanzada": "Contenido de la sección 4.2.3",
                    "4.2.4 Evolución del Comercio Electrónico y Nuevos Modelos de Negocio Digitales": "Contenido de la sección 4.2.4",
                    "4.2.5 Brechas Tecnológicas y Desafíos para la Adopción Digital": "Contenido de la sección 4.2.5"
                },
                "4.3 Análisis de las 5 Fuerzas de Porter": {
                    "4.3.1 Rivalidad entre Competidores": "Contenido de la sección 4.3.1",
                    "4.3.2 Poder de Negociación de los Proveedores": "Contenido de la sección 4.3.2",
                    "4.3.3 Poder de Negociación de los Clientes": "Contenido de la sección 4.3.3",
                    "4.3.4 Amenaza de Nuevos Entrantes": "Contenido de la sección 4.3.4",
                    "4.3.5 Amenaza de Productos Sustitutos": "Contenido de la sección 4.3.5"
                },
                "4.4 Análisis FODA del Sector Empresarial": {
                    "4.4.1 Fortalezas": "Contenido de la sección 4.4.1",
                    "4.4.2 Oportunidades": "Contenido de la sección 4.4.2",
                    "4.4.3 Debilidades": "Contenido de la sección 4.4.3",
                    "4.4.4 Amenazas": "Contenido de la sección 4.4.4"
                },
                "4.5 Mapa de Stakeholders (Actores Clave en el Ecosistema Empresarial)": {
                    "4.5.1 Grandes Empresas": "Contenido de la sección 4.5.1",
                    "4.5.2 PYMEs y Microempresas": "Contenido de la sección 4.5.2",
                    "4.5.3 Startups e Incubadoras de Negocio": "Contenido de la sección 4.5.3",
                    "4.5.4 Cámaras Empresariales y Asociaciones Comerciales": "Contenido de la sección 4.5.4",
                    "4.5.5 Instituciones Gubernamentales y Políticas de Apoyo": "Contenido de la sección 4.5.5",
                    "4.5.6 Universidades, Centros de Investigación e Innovación": "Contenido de la sección 4.5.6",
                    "4.5.7 Organismos Internacionales y ONG con Impacto en la Industria": "Contenido de la sección 4.5.7"
                },
                "4.6 Análisis de Brechas de Competencias (Skill Gap Analysis)": {
                    "4.6.1 Habilidades Técnicas más Demandadas por el Sector Empresarial": "Contenido de la sección 4.6.1",
                    "4.6.2 Habilidades Blandas y de Liderazgo Requeridas en el Mercado Laboral": "Contenido de la sección 4.6.2",
                    "4.6.3 Diferencias entre la Formación Universitaria y las Necesidades Empresariales": "Contenido de la sección 4.6.3",
                    "4.6.4 Estrategias para Alinear la Oferta Académica con la Demanda del Sector": "Contenido de la sección 4.6.4"
                },
                "4.7 Benchmarking del Desarrollo Empresarial": {
                    "4.7.1 Comparación de Cochabamba con otras Ciudades Bolivianas": "Contenido de la sección 4.7.1",
                    "4.7.2 Análisis de Buenas Prácticas en Latinoamérica": "Contenido de la sección 4.7.2",
                    "4.7.3 Estrategias para Mejorar la Competitividad Empresarial en Cochabamba": "Contenido de la sección 4.7.3"
                },
                "4.8 Matriz de Necesidades del Sector Empresarial": {
                    "4.8.1 Talento Humano (Formación y Capacitación)": "Contenido de la sección 4.8.1",
                    "4.8.2 Innovación y Tecnología (Transformación Digital, I+D)": "Contenido de la sección 4.8.2",
                    "4.8.3 Gestión y Financiamiento (Acceso a Créditos, Optimización de Costos)": "Contenido de la sección 4.8.3"
                },
                "4.9 Casos de Éxito de Interacción Academia-Empresa": {
                    "4.9.1 Modelos Exitosos de Universidades en Latinoamérica": "Contenido de la sección 4.9.1",
                    "4.9.2 Empresas Cochabambinas con Colaboración Académica": "Contenido de la sección 4.9.2",
                    "4.9.3 Lecciones Aprendidas y Mejores Prácticas": "Contenido de la sección 4.9.3"
                }
            },
            "5. Estrategias de Interacción Académica – Empresarial": {
                "5.1  Programas de Pasantías y Prácticas Profesionales": "Contenido de la sección 5.1",
                "5.2  Investigación Aplicada y Desarrollo Empresarial": "Contenido de la sección 5.2",
                "5.3  Formación Continua y Capacitación Empresarial": "Contenido de la sección 5.3",
                "5.4  Espacios de Networking y Vinculación": "Contenido de la sección 5.4",
                "5.5  Fomento del Emprendimiento e Innovación Social": "Contenido de la sección 5.5",
                "5.6  Transferencia de Conocimiento y Tecnología": "Contenido de la sección 5.6"
            },
            "6. Beneficios de la Interacción Académica – Empresarial": {
                "6.1 Para la Universidad y la Carrera": "Contenido de la sección 6.1",
                "6.2 Para los Estudiantes": "Contenido de la sección 6.2",
                "6.3 Para las Empresas": "Contenido de la sección 6.3"
            },
            "7. Riesgos y Estrategias de Mitigación": {
                "7.1 Falta de Compromiso de las Empresas": "Contenido de la sección 7.1",
                "7.2 Baja Adaptabilidad de los Estudiantes a las Exigencias del Mercado": "Contenido de la sección 7.2",
                "7.3 Resistencia de la Academia a Cambios en Planes de Estudio": "Contenido de la sección 7.3",
                "7.4 Dificultades en Financiamiento para Proyectos Conjuntos": "Contenido de la sección 7.4"
            },
            "8. Plan de Acción y Cronograma": {
                "8.1 Fase 1: Análisis y Diagnóstico (0-3 meses)": "Contenido de la sección 8.1",
                "8.2 Fase 2: Implementación Piloto (4-12 meses)": "Contenido de la sección 8.2",
                "8.3 Fase 3: Expansión y Evaluación (1-3 años)": "Contenido de la sección 8.3"
            },
            "9. Estrategias de Sostenibilidad y Crecimiento del Proyecto": {
                "9.1 Creación de una Unidad de Vinculación Universidad-Empresa": "Contenido de la sección 9.1",
                "9.2 Modelo de Financiamiento y Continuidad": "Contenido de la sección 9.2",
                "9.3 Evaluación y Actualización del Modelo de Interacción": "Contenido de la sección 9.3"
            },
            "10. Evaluación y Seguimiento": {
                "10.1 Indicadores de Éxito": "Contenido de la sección 10.1"
            },
            "11. Conclusión": {
                "11.1 Reafirmación del Compromiso de la UCB": "Contenido de la sección 11.1",
                "11.2 Invitación a las Empresas a Colaborar Activamente": "Contenido de la sección 11.2",
                "11.3 Sinergia entre la Academia y el Sector Productivo como Motor de Desarrollo": "Contenido de la sección 11.3"
            }
        }
    }


    # Especificar la ruta de salida para el documento Word
    path_out = r"C:\Users\HP\Desktop\Postulacion DOCENCIA\POSTULACION CATO\POSTULACION DOCENTE TIMEPO COMPLETO\Propuesta_Interaccion_Academica_Empresarial.docx"
    # (Opcional) Especificar la ruta del PDF de estilo
    ruta_pdf_estilo = r"C:\Users\HP\Desktop\LIBROS PERSO\CONTEXTO ESPANIOL\CONTEXTO.pdf"

    generar_resumenes_desde_json(json_data, path_out, ruta_pdf_estilo)


# In[ ]:




